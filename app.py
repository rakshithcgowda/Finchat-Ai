import os
import io
import json
import os
import psycopg2
from psycopg2.extras import RealDictCursor
import time
import bcrypt
import docx
from docx.shared import Pt
from fpdf import FPDF
import google.generativeai as genai
from PyPDF2 import PdfReader
import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
from typing import List, Dict
from mistralai import Mistral
import requests
import re
from openai import OpenAI
from PIL import Image
import pdfplumber
import torch
from torchvision import transforms
import uuid
import base64
import tempfile
from docx import Document
import logging
from contextlib import contextmanager

# ─── Configuration ──────────────────────────────────────────────────────────────
GOOGLE_API_KEY = os.environ.get(
    "GOOGLE_API_KEY",
    "AIzaSyANbVVzZACnYnus00xwwRRE01n34yoAmcU"  # fallback for dev/testing
)

MISTRAL_API_KEY = os.environ.get("MISTRAL_API_KEY", "DUW9f3nvZaNkEbxcrxYP4hLIrC3g7Y")
MISTRAL_ENDPOINT = "https://api.mistral.ai/v1/chat/completions"

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "sk-61f7f17d33bd4598b4dd61edd13af337")
DEEPSEEK_CLIENT = OpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com/v1"
)

OCR_API_KEY = "PDF8V7778Y0TX"

BRAND_COLORS = {
    "primary": "#2E86AB",
    "secondary": "#F18F01",
    "background": "#F7F7F7",
    "text": "#121111"
}

genai.configure(api_key=GOOGLE_API_KEY)
mistral_client = Mistral(api_key=MISTRAL_API_KEY)

# ─── OCR.space API Functions ───────────────────────────────────────────────────
OCR_ENDPOINTS = [
    "https://apipro1.ocr.space",
    "https://apipro2.ocr.space",
]

def ocr_space_file_pro(
    filename: str,
    api_key: str = OCR_API_KEY,
    language: str = "eng",
    overlay: bool = False,
    engine: int = 2,
    create_searchable_pdf: bool = True,
    hide_text_layer: bool = True,
    retries: int = 3,
    timeout: int = 120,
) -> dict:
    """Enhanced OCR function with better error handling"""
    mimetype = {
        ".pdf": "application/pdf",
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png"
    }.get(os.path.splitext(filename)[1].lower(), "application/octet-stream")

    payload = {
        "apikey": api_key,
        "language": language,
        "isOverlayRequired": overlay,
        "scale": True,
        "OCREngine": engine,
        "isCreateSearchablePdf": create_searchable_pdf,
        "isSearchablePdfHideTextLayer": hide_text_layer,
    }

    last_err = None
    for endpoint in OCR_ENDPOINTS:
        for attempt in range(1, retries + 1):
            try:
                with open(filename, "rb") as f:
                    files = {"file": (os.path.basename(filename), f, mimetype)}
                    resp = requests.post(
                        endpoint + "/parse/image",
                        data=payload,
                        files=files,
                        timeout=timeout
                    )
                resp.raise_for_status()
                result = resp.json()
                if result.get("IsErroredOnProcessing"):
                    raise RuntimeError(result.get("ErrorMessage") or "OCR processing error")
                return result
            except Exception as e:
                last_err = e
                sleep = 2 ** attempt
                time.sleep(sleep)
    raise RuntimeError(f"All OCR endpoints failed: {last_err}")

def ocr_space_url(
    url: str,
    api_key: str = OCR_API_KEY,
    language: str = "eng",
    overlay: bool = False,
    engine: int = 2,
    retries: int = 3,
    timeout: int = 60,
) -> dict:
    """Enhanced URL OCR function with better error handling"""
    if not url.lower().startswith(("http://", "https://")):
        raise ValueError("URL must start with http:// or https://")

    params = {
        "apikey": api_key,
        "url": url,
        "language": language,
        "isOverlayRequired": overlay,
        "detectOrientation": True,
        "scale": True,
        "OCREngine": engine
    }

    last_err = None
    for endpoint in OCR_ENDPOINTS:
        for attempt in range(1, retries + 1):
            try:
                resp = requests.get(
                    endpoint + "/parse/imageurl",
                    params=params,
                    timeout=timeout
                )
                resp.raise_for_status()
                result = resp.json()
                if result.get("IsErroredOnProcessing"):
                    raise RuntimeError(result.get("ErrorMessage") or "OCR processing error")
                return result
            except Exception as e:
                last_err = e
                time.sleep(2 ** attempt)
    raise RuntimeError(f"All OCR endpoints failed: {last_err}")

# ─── Database Helpers ──────────────────────────────────────────────────────────
DATABASE_URL = 'postgresql://finchat_owner:npg_k0AWSXHr6aqE@ep-ancient-poetry-a6epa959-pooler.us-west-2.aws.neon.tech/finchat?sslmode=require'

DB_CONFIG = {
    'host': 'ep-ancient-poetry-a6epa959-pooler.us-west-2.aws.neon.tech',
    'database': 'finchat',
    'user': 'finchat_owner',
    'password': 'npg_k0AWSXHr6aqE',
    'port': '5432'
}

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

@contextmanager
def get_db_connection():
    conn = None
    try:
        if DATABASE_URL:
            conn = psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)
        else:
            conn = psycopg2.connect(**DB_CONFIG, cursor_factory=RealDictCursor)
        yield conn
    except psycopg2.Error as e:
        logging.error(f"Database connection error: {e}")
        if conn:
            conn.rollback()
        raise
    finally:
        if conn:
            conn.close()

def init_db():
    """Initialize database with enhanced error handling"""
    try:
        with get_db_connection() as conn:
            cursor = conn.cursor()

            # Create users table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS users (
                    username VARCHAR(255) PRIMARY KEY,
                    password BYTEA NOT NULL,
                    role VARCHAR(50) NOT NULL,
                    location_id VARCHAR(255),
                    last_login TIMESTAMP,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Check and add missing columns
            cursor.execute("""
                SELECT column_name
                FROM information_schema.columns
                WHERE table_name = 'users'
            """)
            existing_cols = {row['column_name'] for row in cursor.fetchall()}

            for col, col_type in [
                ("location_id", "VARCHAR(255)"),
                ("last_login", "TIMESTAMP"),
                ("created_at", "TIMESTAMP DEFAULT CURRENT_TIMESTAMP")
            ]:
                if col not in existing_cols:
                    try:
                        cursor.execute(f"ALTER TABLE users ADD COLUMN {col} {col_type}")
                        logging.info(f"Added {col} column to users table")
                    except psycopg2.Error as e:
                        logging.warning(f"Could not add {col} column: {e}")

            # Create subscriptions table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS subscriptions (
                    username VARCHAR(255) PRIMARY KEY,
                    lease_analysis INTEGER DEFAULT 0,
                    deal_structuring INTEGER DEFAULT 0,
                    offer_generator INTEGER DEFAULT 0,
                    FOREIGN KEY(username) REFERENCES users(username) ON DELETE CASCADE
                )
            """)

            # Create interactions table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS interactions (
                    id SERIAL PRIMARY KEY,
                    username VARCHAR(255),
                    feature VARCHAR(100),
                    input_text TEXT,
                    output_text TEXT,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY(username) REFERENCES users(username) ON DELETE CASCADE
                )
            """)

            # Create prompts table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS prompts (
                    id SERIAL PRIMARY KEY,
                    feature VARCHAR(100) UNIQUE,
                    system_prompt TEXT NOT NULL,
                    user_prompt_template TEXT NOT NULL,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            conn.commit()

            # Check if we need to create default data
            cursor.execute("SELECT COUNT(*) FROM users")
            user_count = cursor.fetchone()['count']

            if user_count == 0:
                create_default_admin(conn)
                initialize_default_prompts(conn)

    except psycopg2.Error as e:
        logging.error(f"Database initialization failed: {e}")
        st.error(f"Failed to initialize database: {e}")
        raise

def initialize_default_prompts(conn):
    """Initialize default prompts with better structure"""
    default_prompts = {
        "lease_analysis": {
            "system_prompt": "You are a real estate document expert. Analyze the provided lease agreement and provide a comprehensive summary, including key terms and potential risks.",
            "user_prompt_template": "Summarize this lease agreement in clear, concise language, preserving all key details:\n\n{text}"
        },
        "deal_strategy": {
            "system_prompt": "You are a creative real estate strategist. Based on the provided deal details, suggest structuring options with pros, cons, and negotiation tactics.",
            "user_prompt_template": "Generate creative deal structuring strategies for this real estate deal:\n\n{details}"
        },
        "offer_generator": {
            "system_prompt": "You are a real estate transaction specialist. Generate a professional purchase offer with all essential clauses formatted for the jurisdiction.",
            "user_prompt_template": "Generate a purchase offer based on these details:\n\n{details}"
        }
    }

    cursor = conn.cursor()
    for feature, prompts in default_prompts.items():
        cursor.execute("SELECT 1 FROM prompts WHERE feature = %s", (feature,))
        if not cursor.fetchone():
            cursor.execute(
                "INSERT INTO prompts (feature, system_prompt, user_prompt_template) VALUES (%s, %s, %s)",
                (feature, prompts["system_prompt"], prompts["user_prompt_template"])
            )
    conn.commit()

def create_default_admin(conn):
    """Create default admin with secure password"""
    try:
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM users")
        if cursor.fetchone()['count'] == 0:
            admin_pwd = bcrypt.hashpw("admin123".encode(), bcrypt.gensalt())
            cursor.execute(
                "INSERT INTO users (username, password, role) VALUES (%s, %s, %s)",
                ("admin", admin_pwd, "admin")
            )
            cursor.execute(
                "INSERT INTO subscriptions (username, lease_analysis, deal_structuring, offer_generator) VALUES (%s, %s, %s, %s)",
                ("admin", 1, 1, 1)
            )
            conn.commit()
            logging.info("Default admin user created")
    except psycopg2.Error as e:
        logging.error(f"Failed to create default admin: {e}")
        raise

def verify_password(hashed, password: str) -> bool:
    # Handle different data types that PostgreSQL might return for BYTEA fields
    if isinstance(hashed, memoryview):
        hashed = hashed.tobytes()
    elif isinstance(hashed, str):
        hashed = hashed.encode('utf-8')
    return bcrypt.checkpw(password.encode(), hashed)

# ─── AI Helper Functions ───────────────────────────────────────────────────────
def get_feature_prompt(conn, feature: str, input_text: str) -> tuple:
    """Get prompts with better error handling"""
    cursor = conn.cursor()
    cursor.execute(
        "SELECT system_prompt, user_prompt_template FROM prompts WHERE feature = %s",
        (feature,)
    )
    prompt = cursor.fetchone()
    if prompt:
        system_prompt = prompt['system_prompt']
        user_prompt = prompt['user_prompt_template'].format(text=input_text)
        return system_prompt, user_prompt
    return None, None

def call_gemini(
    conn,
    feature: str,
    content: str,
    temperature: float = 0.7
) -> str:
    """Enhanced Gemini call with better error handling"""
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        system_prompt, user_prompt = get_feature_prompt(conn, feature, content)
        if not system_prompt:
            system_prompt = "You are a knowledgeable assistant."
        prompt = f"SYSTEM: {system_prompt}\n\nUSER: {user_prompt}"
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=temperature,
                top_p=0.95,
                top_k=40,
                max_output_tokens=8192
            )
        )
        return response.text
    except Exception as e:
        st.error(f"Gemini API error: {e}")
        return f"Error: {e}"

def call_deepseek(
    messages: List[Dict[str, str]],
    temperature: float = 0.2,
    top_p: float = 1.0,
    max_tokens: int = 1024,
    stop: List[str] = None,
    stream: bool = False,
    user: str = None,
    logit_bias: Dict[int, float] = None,
) -> str:
    """Enhanced Mistral call with better error handling"""
    payload = {
        "model": "mistral-small-latest",
        "messages": messages,
        "temperature": temperature,
        "top_p": top_p,
        "max_tokens": max_tokens,
    }
    if stop:
        payload["stop"] = stop
    if user:
        payload["user"] = user
    if logit_bias:
        payload["logit_bias"] = logit_bias

    headers = {"Authorization": f"Bearer {MISTRAL_API_KEY}", "Content-Type": "application/json"}
    try:
        resp = requests.post(MISTRAL_ENDPOINT, json=payload, headers=headers, stream=stream)
        resp.raise_for_status()
        data = resp.json()

        if stream:
            return "".join(chunk.get("content", "") for chunk in data.get("choices", []))
        return data["choices"][0]["message"]["content"]
    except Exception as e:
        st.error(f"Mistral API error: {e}")
        return f"Error: {e}"

def call_deepseek(
    messages: List[Dict[str, str]],
    model: str = "deepseek-chat",
    temperature: float = 0.7,
    max_tokens: int = 2000,
    top_p: float = 1.0,
    frequency_penalty: float = 0.0,
    presence_penalty: float = 0.0,
    stream: bool = False,
) -> str:
    """Enhanced DeepSeek call with better error handling"""
    try:
        resp = DEEPSEEK_CLIENT.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=top_p,
            frequency_penalty=frequency_penalty,
            presence_penalty=presence_penalty,
            stream=stream
        )
        if stream:
            return "".join(chunk.choices[0].delta.content for chunk in resp)
        return resp.choices[0].message.content
    except Exception as e:
        return f"Error processing request with DeepSeek: {str(e)}"

def save_interaction(conn, feature: str, input_text: str, output_text: str):
    """Save interaction with better error handling"""
    if st.session_state.get("username"):
        try:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO interactions (username, feature, input_text, output_text) VALUES (%s, %s, %s, %s)",
                (st.session_state.username, feature, input_text, output_text),
            )
            conn.commit()
        except psycopg2.Error as e:
            logging.error(f"Failed to save interaction: {e}")

def extract_text_with_ocr(uploaded_file=None, file_type: str = "pdf", url: str = None):
    """Enhanced text extraction with OCR.space API for PDFs and images."""
    def validate_extracted_text(text: str) -> bool:
        text = text.strip()
        return bool(text) and len(text.split()) >= 2

    def clean_text(text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\x20-\x7E\n]', '', text)
        return text.strip()

    try:
        pages_text = []
        used_ocr = False

        # Handle uploaded file case
        if uploaded_file:
            if file_type.lower() == "pdf":
                # Force OCR for all pages
                with tempfile.TemporaryDirectory() as temp_dir:
                    uploaded_file.seek(0)
                    images = convert_pdf_to_images(uploaded_file, temp_dir, dpi=300)
                    if not images:
                        logging.error("Failed to convert PDF to images for OCR.")
                        return []

                    progress = st.progress(0)
                    ocr_pages = []
                    for i, img_path in enumerate(images):
                        try:
                            img = Image.open(img_path)
                            img = preprocess_image_for_ocr(img)
                            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmpf:
                                img.save(tmpf.name, "PNG", quality=100)
                                resp = ocr_space_file_pro(
                                    filename=tmpf.name,
                                    api_key=OCR_API_KEY,
                                    language='eng',
                                    retries=2,
                                    timeout=60
                                )
                            parsed = resp.get("ParsedResults", [{}])[0].get("ParsedText", "")
                            ocr_pages.append(clean_text(parsed))
                            used_ocr = True
                            os.unlink(tmpf.name)
                        except Exception as e:
                            logging.error(f"OCR failed for page {i+1}: {e}")
                            ocr_pages.append("")
                        progress.progress((i + 1) / len(images))
                        try:
                            os.unlink(img_path)
                        except:
                            pass

                    pages_text = ocr_pages

            # Image processing (JPG)
            else:
                st.info("Processing uploaded JPG with OCR.space API...")
                logging.info("Processing uploaded JPG")
                uploaded_file.seek(0)
                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as temp_file:
                    img = Image.open(uploaded_file)
                    img = preprocess_image_for_ocr(img)
                    img.save(temp_file.name, quality=95)
                    response = ocr_space_file_pro(
                        filename=temp_file.name,
                        api_key=OCR_API_KEY,
                        language='eng',
                        retries=2,
                        timeout=60
                    )
                    os.unlink(temp_file.name)
                parsed_text = response.get("ParsedResults", [{}])[0].get("ParsedText", "")
                cleaned_text = clean_text(parsed_text)
                if validate_extracted_text(cleaned_text):
                    pages_text.append(cleaned_text)
                    used_ocr = True
                    logging.info(f"OCR for JPG: {len(cleaned_text)} chars")
                else:
                    pages_text.append("")
                    st.error("OCR failed to extract meaningful text from JPG.")
                    logging.error("No meaningful text from JPG")

        # Handle URL case (JPG only)
        elif url:
            st.info("Processing JPG URL with OCR...")
            logging.info(f"Processing JPG from URL: {url}")
            response = ocr_space_url(
                url=url,
                api_key=OCR_API_KEY,
                language='eng',
                retries=2,
                timeout=60
            )
            parsed_text = response.get("ParsedResults", [{}])[0].get("ParsedText", "")
            cleaned_text = clean_text(parsed_text)
            if validate_extracted_text(cleaned_text):
                pages_text.append(cleaned_text)
                used_ocr = True
                logging.info(f"OCR for URL JPG: {len(cleaned_text)} chars")
            else:
                pages_text.append("")
                st.error("OCR failed to extract meaningful text from JPG URL.")
                logging.error("No meaningful text from JPG URL")

        # Validate final results
        if not any(validate_extracted_text(p) for p in pages_text):
            st.error("Failed to extract meaningful text from the document.")
            logging.error("No meaningful text extracted from document")
            return []

        st.session_state['used_ocr'] = used_ocr
        logging.info(f"Extraction complete: {len(pages_text)} pages, used_ocr={used_ocr}")
        return pages_text

    except Exception as e:
        logging.error(f"extract_text_with_ocr failed: {e}", exc_info=True)
        st.error(f"Text extraction failed: {e}")
        return []

def convert_pdf_to_images(pdf_file, output_dir, dpi=300):
    """Convert PDF to high-quality images for OCR with better error handling"""
    images = []
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for i, page in enumerate(pdf.pages):
                img = page.to_image(resolution=dpi).original
                img_path = os.path.join(output_dir, f"page_{i+1}.png")
                img.save(img_path, "PNG", quality=100)
                images.append(img_path)
    except Exception as e:
        st.error(f"PDF to image conversion failed: {e}")
    return images

def preprocess_image_for_ocr(img):
    """Enhanced image preprocessing for better OCR results"""
    try:
        # Convert to grayscale
        if img.mode != 'L':
            img = img.convert('L')

        # Enhance contrast
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)

        # Remove noise
        img = img.point(lambda x: 0 if x < 128 else 255, '1')

        return img
    except Exception as e:
        st.warning(f"Image preprocessing failed: {e}")
        return img

import re
from fpdf import FPDF
import streamlit as st
from datetime import datetime
import os

def lease_summarization_ui(conn):
    """Simplified LeaseBrief Buddy with single-button workflow"""
    st.header("📄 LeaseBrief Buddy")

    # Clear previous summary and related data
    if st.button("Clear Summary & Chat", key="clear_lease_summary"):
        for k in ['last_file', 'last_url', 'last_summary', 'extracted_pages', 'lease_chat_memory']:
            st.session_state.pop(k, None)
        st.success("Cleared previous summary")
        st.rerun()

    st.markdown("Upload your lease PDF or JPG image to get an AI-generated summary.")

    # File uploader for PDF or JPG
    uploaded_file = st.file_uploader(
        "Upload Lease Document (PDF or JPG)",
        type=["pdf", "jpg", "jpeg"],
        key="lease_file_uploader"
    )

    # URL input for remote JPG
    image_url = st.text_input(
        "Or Enter Public JPG URL",
        key="lease_image_url",
        placeholder="e.g., https://example.com/image.jpg"
    )

    # Validate inputs
    if uploaded_file and image_url:
        st.error("Please provide either a file or a URL, not both.")
        return
    elif not uploaded_file and not image_url:
        st.info("Please upload a file or enter a URL to proceed.")
        return

    # Generate summary button - combines OCR and summarization
    if st.button("Generate Summary", key="lease_generate_button"):
        with st.spinner("Processing document..."):
            try:
                if uploaded_file:
                    file_type = "pdf" if uploaded_file.name.lower().endswith(".pdf") else "jpg"
                    pages = extract_text_with_ocr(uploaded_file=uploaded_file, file_type=file_type)
                    text = "\n".join(pages) if pages else ""
                    input_identifier = uploaded_file.name
                elif image_url:
                    pages = extract_text_with_ocr(url=image_url)
                    text = "\n".join(pages) if pages else ""
                    input_identifier = image_url

                if not text.strip():
                    st.error("Failed to extract meaningful text from the document.")
                    return

                # Generate summary
                summary = call_deepseek(
                    messages=[{"role": "user", "content": f"Summarize this lease agreement in clear, concise language, highlighting key terms and potential issues:\n\n{text}"}],
                    temperature=0.3,
                    max_tokens=2000
                )

                # Store results
                st.session_state.last_summary = summary
                st.session_state.last_file = input_identifier
                st.session_state.extracted_pages = pages if pages else []
                save_interaction(conn, "lease_summary_full", input_identifier, summary)
                st.rerun()

            except Exception as e:
                st.error(f"Error processing document: {str(e)}")
                return

    # Display existing summary if available
    if 'last_summary' in st.session_state:
        st.subheader("Document Summary")
        st.markdown(st.session_state['last_summary'])

        # Export section
        st.markdown("### 📥 Export Summary")

        # Derive default filename
        file_base = os.path.splitext(st.session_state.get('last_file', 'lease_summary'))[0]
        file_name = st.text_input("Filename (no extension):", value=file_base, key="lease_export_name")

        # Function to clean markdown from text
        def clean_markdown(text):
            # Remove markdown headers (e.g., ###, ##, #)
            text = re.sub(r'^#+ ', '', text, flags=re.MULTILINE)
            # Remove bold/italic markers (e.g., **text**, *text*)
            text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^\*]+)\*', r'\1', text)
            # Remove list markers (e.g., - or *)
            text = re.sub(r'^\s*[-*]\s+', '', text, flags=re.MULTILINE)
            # Remove extra newlines and leading/trailing whitespace
            text = re.sub(r'\n\s*\n', '\n', text)
            text = text.strip()
            return text

        # PDF Export
        class LeasePDF(FPDF):
            def header(self):
                self.set_font('Arial', 'B', 20)
                self.set_text_color(30, 90, 140)
                self.cell(0, 12, "Lease Agreement Summary", ln=1, align='C')
                self.set_font('Arial', 'I', 10)
                self.set_text_color(120, 120, 120)
                date_str = datetime.now().strftime("%B %d, %Y")
                self.cell(0, 6, date_str, ln=1, align='C')
                self.set_draw_color(200, 200, 200)
                self.set_line_width(0.4)
                y = self.get_y() + 2
                self.line(15, y, 195, y)
                self.ln(8)

            def footer(self):
                self.set_y(-15)
                self.set_font('Arial', 'I', 8)
                self.set_text_color(130, 130, 130)
                page_info = f"Page {self.page_no()} / {{nb}}"
                self.cell(0, 10, page_info, align='C')

            def section_title(self, title):
                self.set_font('Arial', 'B', 14)
                self.set_fill_color(245, 245, 245)
                self.set_text_color(30, 90, 140)
                self.cell(0, 8, title, ln=1, fill=True, border=0)
                x1, x2 = self.l_margin, self.w - self.r_margin
                y = self.get_y() - 1
                self.set_draw_color(30, 90, 140)
                self.set_line_width(0.3)
                self.line(x1, y, x2, y)
                self.ln(4)

            def paragraph(self, text):
                self.set_font('Arial', '', 12)
                self.set_text_color(50, 50, 50)
                # Encode text to handle special characters
                self.multi_cell(0, 7, text.encode('latin-1', 'replace').decode('latin-1'))
                self.ln(4)

        pdf = LeasePDF()
        pdf.set_auto_page_break(True, margin=20)
        pdf.add_page()
        pdf.set_left_margin(20)
        pdf.set_right_margin(20)
        pdf.set_title(f"Lease Summary - {file_name}")
        pdf.set_author("Property Deals AI")

        pdf.section_title("Document Information")
        info = [
            ("Input File", st.session_state.get('last_file', 'Unknown')),
            ("Generated at", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ]
        for i, (label, val) in enumerate(info):
            pdf.set_fill_color(245, 245, 245) if i % 2 else pdf.set_fill_color(255, 255, 255)
            pdf.set_font('Arial', 'B', 11)
            pdf.cell(50, 6, f"{label}:", ln=0, fill=True)
            pdf.set_font('Arial', '', 11)
            pdf.cell(0, 6, str(val).encode('latin-1', 'replace').decode('latin-1'), ln=1, fill=True)
        pdf.ln(6)

        pdf.section_title("Lease Summary")
        # Clean the markdown from the summary before adding to PDF
        cleaned_summary = clean_markdown(st.session_state['last_summary'])
        pdf.paragraph(cleaned_summary)

        try:
            pdf_bytes = pdf.output(dest='S').encode('latin-1', 'ignore')
            st.download_button(
                "Download PDF",
                pdf_bytes,
                file_name=f"{file_name}.pdf",
                mime="application/pdf",
                key="lease_export_pdf"
            )
        except Exception as e:
            st.error(f"PDF generation failed: {e}")

        # Chat functionality
        st.divider()

        # Initialize chat memory if not exists
        if 'lease_chat_memory' not in st.session_state:
            st.session_state.lease_chat_memory = []
            st.session_state.lease_chat_memory.clear()

        # Display chat history
        for role, message in st.session_state.lease_chat_memory:
            st.chat_message(role).write(message)

        # Display the header Chatbot
        st.subheader("Chatbot")
        # Smaller header: You can question based on the above summary
        st.markdown("Ask questions based on the lease summary.")
        # Chat input
        user_input = st.chat_input()
        if user_input:
            # Add user message to chat
            st.session_state.lease_chat_memory.append(("user", user_input))
            st.chat_message("user").write(user_input)

            # Build context from the lease summary
            context = f"Lease Summary:\n{st.session_state['last_summary']}\n\nQuestion: {user_input}"

            # Call AI
            with st.spinner("Generating response..."):
                try:
                    response = call_deepseek(
                        messages=[
                            {"role": "system", "content": "You are a lease agreement expert. Answer questions based on the provided lease summary."},
                            {"role": "user", "content": context}
                        ],
                        temperature=0.2
                    )

                    # Add AI response to chat
                    st.session_state.lease_chat_memory.append(("assistant", response))
                    st.chat_message("assistant").write(response)

                    # Save interaction
                    save_interaction(conn, "lease_chat", user_input, response)
                except Exception as e:
                    st.error(f"Chat response failed: {e}")

def deal_structuring_ui(conn):
    """Property Guru - AI-powered deal strategist"""
    st.header("💡 Property Guru")
    st.markdown("Get AI-powered strategies tailored to your property deal.")

    # Initialize session state
    if "property_guru_memory" not in st.session_state:
        st.session_state.property_guru_memory = []
        st.session_state.last_strategies = None
        st.session_state.strategy_chat_memory = []

    # Clear chat and strategies
    if st.button("Clear Strategies & Chat"):
        st.session_state.property_guru_memory.clear()
        st.session_state.last_strategies = None
        st.session_state.strategy_chat_memory.clear()
        st.rerun()

    # Input form for Buyer and Seller details
    with st.expander("Deal Details", expanded=True):
        st.markdown("### Buyer Situation")
        col1, col2 = st.columns(2)
        with col1:
            deposit = st.selectbox("Available Deposit/Upfront Cash",
                                 ["Low/Zero", "Moderate", "High"],
                                 key="buyer_deposit")
            credit = st.selectbox("Credit & Mortgage Ability",
                                ["Excellent", "Average", "Poor", "None"],
                                key="buyer_credit")
            experience = st.selectbox("Property Investing Experience",
                                    ["No creative strategy done yet",
                                     "Have done at least one"],
                                    key="buyer_experience")
        with col2:
            risk_appetite = st.selectbox("Risk Appetite",
                                       ["Low – very cautious",
                                        "Medium – balanced",
                                        "High – very comfortable with risk"],
                                       key="buyer_risk")
            goal = st.selectbox("Primary Goal",
                              ["Hold long-term (rental/residence)",
                               "Resell for profit ASAP"],
                              key="buyer_goal")
            timeline = st.selectbox("Investment Timeline",
                                  ["Few months", "1-2 years", "Longer period"],
                                  key="buyer_timeline")

        st.markdown("### Property Details")
        col1, col2 = st.columns(2)
        with col1:
            property_type = st.selectbox("Property Type",
                                       ["Residential", "Commercial",
                                        "Mixed-Use", "Land"],
                                       key="property_type")
            market_price = st.number_input("Market Price (£)",
                                         min_value=0, step=1000,
                                         key="market_price")
            offer_price = st.number_input("Your Offer Price (£)",
                                        min_value=0, step=1000,
                                        key="offer_price")
        with col2:
            property_condition = st.text_area("Property Condition & Repairs Needed",
                                            placeholder="e.g., good condition, needs roof replacement, etc.",
                                            key="property_condition")
            market_condition = st.selectbox("Market Condition",
                                          ["Hot (multiple offers)",
                                           "Moderate (some interest)",
                                           "Slow (little interest)"],
                                          key="market_condition")

        st.markdown("### Seller Motivation & Situation")
        seller_motivation = st.text_area("Seller's Motivation & Urgency",
                                        placeholder="e.g., relocating, financial distress, etc.",
                                        key="seller_motivation")

    # Generate strategies
    if st.button("Generate Strategies", type="primary"):
        prompt = (
            f"Buyer Situation:\n"
            f"- Available Deposit/Upfront Cash: {deposit}\n"
            f"- Credit & Mortgage Ability: {credit}\n"
            f"- Property Investing Experience: {experience}\n"
            f"- Risk Appetite: {risk_appetite}\n"
            f"- Primary Goal: {goal}\n"
            f"- Investment Timeline: {timeline}\n\n"
            f"Property Details:\n"
            f"- Property Type: {property_type}\n"
            f"- Market Price: £{market_price}\n"
            f"- Offer Price: £{offer_price}\n"
            f"- Property Condition: {property_condition}\n"
            f"- Market Condition: {market_condition}\n\n"
            f"Seller Motivation:\n"
            f"- Motivation & Urgency: {seller_motivation}\n\n"
            f"Generate 3 specific deal structuring strategies for this real estate deal. "
            f"Provide clear actionable strategies with implementation steps. "
            f"Format each strategy with a numbered heading and bullet points."
        )

        with st.spinner("Developing strategies..."):
            messages = [
                {"role": "system", "content": "You are a creative real estate strategist. Provide specific actionable deal structuring options with pros and cons."},
                {"role": "user", "content": prompt}
            ]
            strategies = call_deepseek(messages=messages, temperature=0.7)

        # Store strategies
        st.session_state.property_guru_memory.append(("assistant", strategies))
        st.session_state.last_strategies = strategies
        save_interaction(conn, "deal_strategy", prompt, strategies)
        st.rerun()  # Rerun to update UI with strategies

    # Display strategies if they exist
    if st.session_state.get("last_strategies"):
        st.subheader("Recommended Strategies")
        st.markdown(st.session_state.last_strategies)

        # Chat functionality
        st.divider()
        st.subheader("Chatbot")
        st.markdown("Ask questions based on the strategies above.")

        # Display chat history
        for role, message in st.session_state.strategy_chat_memory:
            with st.chat_message(role):
                st.write(message)

        # Chat input
        user_input = st.chat_input("Ask a question about the strategies")
        if user_input:
            # Add user message to chat
            st.session_state.strategy_chat_memory.append(("user", user_input))

            # Build context from the strategies
            context = f"Strategies:\n{st.session_state.last_strategies}\n\nQuestion: {user_input}"

            # Call AI
            with st.spinner("Generating response..."):
                response = call_deepseek(
                    messages=[
                        {"role": "system", "content": "You are a real estate strategy expert answering questions about deal structuring."},
                        {"role": "user", "content": context}
                    ],
                    temperature=0.3
                )

            # Add AI response to chat
            st.session_state.strategy_chat_memory.append(("assistant", response))
            save_interaction(conn, "strategy_chat", user_input, response)
            st.rerun()  # Rerun to update chat UI

          
def build_guided_prompt(details, detail_level):
    """Build a plain text prompt for generating a clean, empathetic offer letter without markdown."""
    
    vendor = details['vendor']
    local_knowledge = details['local_knowledge']
    costs = details['costs']
    negative_points = details['negative_points']
    offers = details['offers']
    social_proof = details['social_proof']
    proof_of_funds = details['proof_of_funds']

    # Prepare sub-components outside the main f-string to avoid syntax issues
    comparables_str = ', '.join([
        f"{comp['address']} ({comp['status']}, £{comp['price']})"
        for comp in local_knowledge.get('comparables', [])
    ])

    offers_str = ', '.join([
        f"£{offer['amount']} ({offer['description']})"
        for offer in offers
    ])

    costs_str = ', '.join([
        f"{k}: £{v}" for k, v in costs.items() if v > 0
    ])

    prompt = (
        "Draft a professional and empathetic offer letter for a real estate transaction. "
        "The letter must be in plain text with no markdown characters (e.g., #, *, **, -, etc.), "
        "no bullet points, and no headers. Use clear, concise paragraphs separated by single newlines. "
        "Address the letter to the vendor by name and acknowledge their situation empathetically. "
        "Include details about the property and local market, referencing comparable sales and average sale price. "
        "Present the proposed offers with their amounts and descriptions in a narrative format, not as a list. "
        "Explain the benefits of the offer (e.g., no estate agent fees, covered legal costs, flexible completion, 15% profit margin). "
        "Justify the 15% profit margin. Suggest next steps, such as preparing a Heads of Terms document. "
        "End with a polite closing and include placeholders for the sender's name, position, and contact information. "
        f"Vendor Name: {vendor.get('name', 'N/A')}. "
        f"Vendor Situation: {vendor.get('situation', 'N/A')}. "
        f"Property Condition: {vendor.get('condition', 'N/A')}. "
        f"Other Challenges: {vendor.get('challenges') or 'None provided'}. "
        f"Family Member: {vendor.get('family_member') or 'None provided'}. "
        f"Family Situation: {vendor.get('family_situation') or 'None provided'}. "
        f"Local Investment Experience: {local_knowledge.get('experience', 'N/A')}. "
        f"Comparable Sales: {comparables_str or 'None provided'}. "
        f"Average Sale Price: £{local_knowledge.get('avg_sale_price', 'N/A')}. "
        f"Required Profit: £{local_knowledge.get('required_profit', 'N/A')}. "
        f"Profit Reason: {local_knowledge.get('profit_reason', 'N/A')}. "
        f"Costs: {costs_str or 'None provided'}. "
        f"Negative Points: {negative_points or 'None provided'}. "
        f"Offers: {offers_str or 'None provided'}. "
        f"Social Proof: {social_proof or 'None provided'}. "
        f"Proof of Funds: {proof_of_funds or 'None provided'}. "
        f"Detail Level: {detail_level} (Minimal: brief letter, Standard: balanced with key details, Comprehensive: include all provided details)."
    )

    return prompt



def offer_generator_ui(conn):
    """Enhanced Offer Buddy with fixed comments and download functionality, producing plain text output"""
    st.header("✍️ Offer Buddy")
    st.markdown(
        """
        <style>
        .offer-section { background-color: #f0f2f6; border-radius: 10px; padding: 15px; margin-bottom: 20px; }
        .offer-highlight { background-color: #fffacd; padding: 2px 5px; border-radius: 3px; }
        </style>
        """, unsafe_allow_html=True
    )

    # Initialize session state for Offer Buddy
    if 'offer_stage' not in st.session_state:
        st.session_state.update({
            'offer_stage': 'details_entry',
            'offer_data': {},
            'generated_offer': None,
            'edited_offer': None,
            'review_comments': [],
            'form_submitted': False
        })

    # Navigation tabs for stages
    stages = ["details_entry", "offer_generation", "review_edit", "export"]
    labels = ["Details Entry", "Offer Generation", "Review & Edit", "Export"]
    idx = stages.index(st.session_state.offer_stage) if st.session_state.offer_stage in stages else 0
    cols = st.columns(len(stages))
    for i, label in enumerate(labels):
        with cols[i]:
            if i < idx:
                st.success(f"✓ {label}")
            elif i == idx:
                st.info(f"→ {label}")
            else:
                st.caption(label)

    # Stage 1: Details Entry
    if st.session_state.offer_stage == 'details_entry':
        st.markdown("### 1. Enter Offer Details")
        with st.form("offer_details_form", clear_on_submit=False):
            st.markdown('<div class="offer-section">', unsafe_allow_html=True)
            st.markdown('#### Vendor Details')
            vendor_name = st.text_input("Vendor Name*", key="vendor_name_input", placeholder="Enter vendor name")
            vendor_situation = st.text_area("Vendor Situation*", key="vendor_situation_input", placeholder="e.g., relocating, financial distress")
            property_condition = st.text_area("Property Condition*", key="property_condition_input", placeholder="e.g., needs roof replacement")
            other_challenges = st.text_area("Other Challenges", key="other_challenges_input", placeholder="e.g., planning issues, structural concerns")
            family_member = st.text_input("Family Member Name", key="family_member_input")
            family_situation = st.text_area("Family Member Situation", key="family_situation_input", placeholder="e.g., elderly parent needing care")

            st.markdown('#### Local Knowledge & Comparables')
            local_experience = st.text_area("Local Investment Experience*", key="local_experience_input", placeholder="e.g., I have been investing in [Local Area] for several years")
            st.markdown("**Recent Comparables**")
            comparables = []
            for i in range(3):
                st.markdown(f"**Comparable {i+1}**")
                col1, col2, col3 = st.columns(3)
                with col1:
                    address = st.text_input(f"Property Address {i+1}*", key=f"comp_address_{i}_input")
                with col2:
                    status = st.selectbox(f"Status {i+1}*", ["Sold", "Under Offer", "SSTC"], key=f"comp_status_{i}_input")
                with col3:
                    price = st.number_input(f"Price {i+1}* (£)", min_value=0.0, step=1000.0, key=f"comp_price_{i}_input")
                comparables.append({"address": address, "status": status, "price": price})
            avg_sale_price = st.number_input("Average Sale Price* (£)", min_value=0.0, step=1000.0, key="avg_sale_price_input")
            required_profit = st.number_input("Required Profit (15%)* (£)", min_value=0.0, step=1000.0, key="required_profit_input")
            profit_reason = st.text_area("Reason for Profit*", value="I maintain at least a 15% profit margin to protect against unforeseen risks.", key="profit_reason_input")

            st.markdown('#### Detailed Costs')
            costs = {
                "Deposit": st.number_input("Deposit (£)", min_value=0.0, step=1000.0, key="cost_deposit_input"),
                "Finance": st.number_input("Finance (£)", min_value=0.0, step=1000.0, key="cost_finance_input"),
                "Solicitors": st.number_input("Solicitors (£)", min_value=0.0, step=500.0, key="cost_solicitors_input"),
                "Stamp Duty": st.number_input("Stamp Duty (£)", min_value=0.0, step=500.0, key="cost_stamp_duty_input"),
                "Estate Agent Fees": st.number_input("Estate Agent Fees (£)", min_value=0.0, step=500.0, key="cost_estate_agent_input"),
                "Doors and Windows": st.number_input("Doors and Windows (£)", min_value=0.0, step=500.0, key="cost_doors_windows_input"),
                "Joists and Floorboards": st.number_input("Joists and Floorboards (£)", min_value=0.0, step=500.0, key="cost_joists_input"),
                "Chimney Breast": st.number_input("Chimney Breast (£)", min_value=0.0, step=500.0, key="cost_chimney_input"),
                "Lintels & Structural Work": st.number_input("Lintels & Structural Work (£)", min_value=0.0, step=500.0, key="cost_lintels_input"),
                "Roof, Guttering & Fascias": st.number_input("Roof, Guttering & Fascias (£)", min_value=0.0, step=500.0, key="cost_roof_input"),
                "External Wall Render": st.number_input("External Wall Render (£)", min_value=0.0, step=500.0, key="cost_render_input"),
                "Electrical Wiring & Lighting": st.number_input("Electrical Wiring & Lighting (£)", min_value=0.0, step=500.0, key="cost_electrical_input"),
                "Wi-Fi Points": st.number_input("Wi-Fi Points (£)", min_value=0.0, step=100.0, key="cost_wifi_input"),
                "Fire Safety Panels": st.number_input("Fire Safety Panels (£)", min_value=0.0, step=500.0, key="cost_fire_safety_input"),
                "Studs / Plaster Boards / Insulation": st.number_input("Studs / Plaster Boards / Insulation (£)", min_value=0.0, step=500.0, key="cost_studs_input"),
                "Emergency Lights": st.number_input("Emergency Lights (£)", min_value=0.0, step=100.0, key="cost_emergency_lights_input"),
                "Cupboards & Skirting Boards": st.number_input("Cupboards & Skirting Boards (£)", min_value=0.0, step=500.0, key="cost_cupboards_input"),
                "Window Sills": st.number_input("Window Sills (£)", min_value=0.0, step=100.0, key="cost_window_sills_input"),
                "Locks & Ironmongery": st.number_input("Locks & Ironmongery (£)", min_value=0.0, step=100.0, key="cost_locks_input"),
                "Plumbing & Heating": st.number_input("Plumbing & Heating (£)", min_value=0.0, step=500.0, key="cost_plumbing_input"),
                "Kitchen": st.number_input("Kitchen (£)", min_value=0.0, step=500.0, key="cost_kitchen_input"),
                "Bathroom": st.number_input("Bathroom (£)", min_value=0.0, step=500.0, key="cost_bathroom_input"),
                "Tiling": st.number_input("Tiling (£)", min_value=0.0, step=500.0, key="cost_tiling_input"),
                "Painting & Decoration": st.number_input("Painting & Decoration (£)", min_value=0.0, step=500.0, key="cost_painting_input"),
                "Flooring / Carpet": st.number_input("Flooring / Carpet (£)", min_value=0.0, step=500.0, key="cost_flooring_input"),
                "Garden & Landscaping": st.number_input("Garden & Landscaping (£)", min_value=0.0, step=500.0, key="cost_garden_input")
            }

            st.markdown('#### Potential Negative Points')
            negative_points = st.text_area("Potential Negative Points", key="negative_points_input", placeholder="e.g., market risks, structural issues")

            st.markdown('#### Proposed Offers')
            offers = []
            for i in range(3):
                st.markdown(f"**Offer {i+1}**")
                col1, col2 = st.columns(2)
                with col1:
                    amount = st.number_input(f"Offer {i+1} Amount (£)*", min_value=0.0, step=1000.0, key=f"offer_amount_{i}_input")
                with col2:
                    description = st.text_input(f"Offer {i+1} Description*", key=f"offer_desc_{i}_input", placeholder="e.g., Cash offer with quick close")
                offers.append({"amount": amount, "description": description})

            st.markdown('#### Social Proof & Proof of Funds')
            social_proof = st.text_input("Google Reviews Link", key="social_proof_input")
            proof_of_funds = st.text_area("Proof of Funds", key="proof_of_funds_input", placeholder="e.g., Bank statement summary")

            st.markdown('#### AI Configuration')
            creativity = st.slider("Creativity Level", 0.0, 1.0, 0.3, key="offer_creativity_input")
            detail_level = st.select_slider(
                "Detail Level", options=["Minimal", "Standard", "Comprehensive"],
                value="Standard", key="offer_detail_level_input"
            )

            st.markdown('</div>', unsafe_allow_html=True)

            submitted = st.form_submit_button("Generate Offer Draft", type="primary")
            if submitted:
                missing = []
                for field, msg in {
                    "vendor_name_input": "Vendor name required",
                    "vendor_situation_input": "Vendor situation required",
                    "property_condition_input": "Property condition required",
                    "local_experience_input": "Local experience required",
                    "avg_sale_price_input": "Average sale price required",
                    "required_profit_input": "Required profit required",
                    "profit_reason_input": "Profit reason required"
                }.items():
                    if not st.session_state.get(field):
                        missing.append(msg)
                for i in range(3):
                    if not st.session_state.get(f"comp_address_{i}_input") or \
                       not st.session_state.get(f"comp_status_{i}_input") or \
                       not st.session_state.get(f"comp_price_{i}_input"):
                        missing.append(f"Comparable {i+1} details required")
                    if not st.session_state.get(f"offer_amount_{i}_input") or \
                       not st.session_state.get(f"offer_desc_{i}_input"):
                        missing.append(f"Offer {i+1} amount and description required")
                if missing:
                    for m in missing:
                        st.error(m)
                else:
                    st.session_state.form_submitted = True
                    st.session_state.offer_data['details'] = {
                        'vendor': {
                            'name': vendor_name,
                            'situation': vendor_situation,
                            'condition': property_condition,
                            'challenges': other_challenges,
                            'family_member': family_member,
                            'family_situation': family_situation
                        },
                        'local_knowledge': {
                            'experience': local_experience,
                            'comparables': comparables,
                            'avg_sale_price': avg_sale_price,
                            'required_profit': required_profit,
                            'profit_reason': profit_reason
                        },
                        'costs': costs,
                        'negative_points': negative_points,
                        'offers': offers,
                        'social_proof': social_proof,
                        'proof_of_funds': proof_of_funds
                    }
                    st.session_state.offer_data.update({
                        'creativity': creativity,
                        'detail_level': detail_level
                    })
                    st.session_state.offer_stage = 'offer_generation'
                    st.rerun()

        # Clear form button
        if st.button("Clear Form", key="clear_offer_form"):
            for key in list(st.session_state.keys()):
                if key.startswith("vendor_") or key.startswith("comp_") or \
                   key.startswith("cost_") or key.startswith("offer_") or \
                   key in ["negative_points_input", "social_proof_input", "proof_of_funds_input",
                           "avg_sale_price_input", "required_profit_input", "profit_reason_input",
                           "local_experience_input", "offer_creativity_input", "offer_detail_level_input"]:
                    del st.session_state[key]
            st.session_state.form_submitted = False
            st.rerun()

    # Stage 2: Offer Generation
    if st.session_state.offer_stage == 'offer_generation':
        d = st.session_state.offer_data
        prompt = build_guided_prompt(d['details'], d['detail_level'])

        with st.spinner("Generating offer draft..."):
            messages = [
                {
                    'role': 'system',
                    'content': (
                        'You are a real estate attorney drafting an empathetic offer letter. '
                        'Produce a professional, plain text letter without any markdown characters '
                        '(e.g., #, *, **, -, etc.), bullet points, or headers. Use clear, concise language '
                        'with paragraphs separated by single newlines. Avoid numbered lists or formatting symbols.'
                    )
                },
                {'role': 'user', 'content': prompt}
            ]
            offer = call_deepseek(messages, temperature=d['creativity'])
            st.session_state.generated_offer = offer
            save_interaction(conn, 'offer_generator', prompt, offer)

        st.subheader("Generated Offer Draft")
        st.text(offer)  # Display as plain text
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Proceed to Review", type="primary"):
                st.session_state.offer_stage = 'review_edit'
                st.rerun()
        with col2:
            if st.button("← Back to Details"):
                st.session_state.offer_stage = 'details_entry'
                st.rerun()

    # Stage 3: Review & Edit
    if st.session_state.offer_stage == 'review_edit':
        st.subheader("Review & Edit Offer")
        edited = st.text_area(
            "Edit Offer Draft",
            value=st.session_state.generated_offer,
            height=400,
            key='offer_edit_input'
        )
        if edited != st.session_state.edited_offer:
            st.session_state.edited_offer = edited

        st.markdown("#### Comments")
        new_comment = st.text_input("Add a Comment", key='offer_new_comment_input')
        if st.button("Add Comment", key="add_comment_button") and new_comment:
            ts = datetime.now().strftime("%Y-%m-%d %H:%M")
            st.session_state.review_comments.append({
                'ts': ts,
                'text': new_comment,
                'resolved': False
            })
            st.rerun()

        for i, comment in enumerate(st.session_state.review_comments):
            cols = st.columns([1, 8, 1])
            with cols[0]:
                st.markdown(f"**{comment['ts']}**")
            with cols[1]:
                st.markdown(f"{'✓' if comment['resolved'] else '◯'} {comment['text']}")
            with cols[2]:
                if not comment['resolved'] and st.button('Resolve', key=f'resolve_comment_{i}'):
                    comment['resolved'] = True
                    st.rerun()

        col1, col2 = st.columns(2)
        with col1:
            if st.button('← Back to Generation'):
                st.session_state.offer_stage = 'offer_generation'
                st.rerun()
        with col2:
            if st.button('Proceed to Export', type="primary"):
                st.session_state.offer_stage = 'export'
                st.rerun()

    # Stage 4: Export
    if st.session_state.offer_stage == 'export':
        st.subheader("Export Offer")
        content = st.session_state.edited_offer or st.session_state.generated_offer
        if st.checkbox('Include Comments in Export', value=True, key='include_comments_input'):
            comments_section = "\n\nComments\n" + \
                "\n".join([f"{c['ts']}: {c['text']}" for c in st.session_state.review_comments])
            content += comments_section

        export_format = st.selectbox(
            'Export Format',
            ['PDF', 'Word', 'Text', 'HTML'],
            key='offer_export_format_input'
        )
        name = st.text_input(
            'File Name (without extension)',
            value='property_offer',
            key='offer_export_name_input'
        )

        if st.button('Download Offer', type="primary", key='download_offer_button'):
            if export_format == 'PDF':
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font('Arial', size=12)
                for line in content.split('\n'):
                    pdf.multi_cell(0, 6, line.encode('latin-1', 'replace').decode('latin-1'))
                pdf_output = pdf.output(dest='S').encode('latin-1')
                st.download_button(
                    'Download PDF',
                    pdf_output,
                    f"{name}.pdf",
                    "application/pdf",
                    key='download_pdf_button'
                )
            elif export_format == 'Word':
                doc = Document()
                for line in content.split('\n'):
                    doc.add_paragraph(line)  # Plain text paragraphs
                buf = io.BytesIO()
                doc.save(buf)
                st.download_button(
                    'Download Word',
                    buf.getvalue(),
                    f"{name}.docx",
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    key='download_word_button'
                )
            elif export_format == 'HTML':
                html = f"<pre>{content}</pre>"
                st.download_button(
                    'Download HTML',
                    html.encode(),
                    f"{name}.html",
                    'text/html',
                    key='download_html_button'
                )
            else:
                st.download_button(
                    'Download Text',
                    content.encode(),
                    f"{name}.txt",
                    'text/plain',
                    key='download_text_button'
                )

        col1, col2 = st.columns(2)
        with col1:
            if st.button('← Back to Review'):
                st.session_state.offer_stage = 'review_edit'
                st.rerun()
        with col2:
            if st.button('Start New Offer', type="primary"):
                for key in list(st.session_state.keys()):
                    if key.startswith('offer_') or key in [
                        'vendor_name_input', 'vendor_situation_input', 'property_condition_input',
                        'other_challenges_input', 'family_member_input', 'family_situation_input',
                        'local_experience_input', 'avg_sale_price_input', 'required_profit_input',
                        'profit_reason_input', 'negative_points_input', 'social_proof_input',
                        'proof_of_funds_input', 'form_submitted'
                    ] or key.startswith('comp_') or key.startswith('cost_') or key.startswith('offer_amount_') or key.startswith('offer_desc_'):
                        del st.session_state[key]
                st.session_state.offer_stage = 'details_entry'
                st.rerun()

def go_buddy_ui(conn):
    """Enhanced GO Buddy with better OCR handling and chat functionality"""
    st.header("🚀 Auction Buddy")
    st.markdown("Upload multiple documents and get individual summaries in one merged document.")

    if 'go_buddy_files' not in st.session_state:
        st.session_state.go_buddy_files = []
        st.session_state.go_buddy_summaries = []
        st.session_state.go_buddy_chat_memory = []

    uploaded_files = st.file_uploader(
        "Upload Documents (PDF or JPG)",
        type=["pdf", "jpg", "jpeg"],
        accept_multiple_files=True,
        key="go_buddy_uploader"
    )

    if uploaded_files and st.button("Generate Summaries"):
        st.session_state.go_buddy_files = uploaded_files
        st.session_state.go_buddy_summaries = []
        st.session_state.go_buddy_chat_memory = []

        progress_bar = st.progress(0)
        for i, uploaded_file in enumerate(uploaded_files):
            with st.spinner(f"Processing {uploaded_file.name} ({i+1}/{len(uploaded_files)})..."):
                file_type = "pdf" if uploaded_file.name.lower().endswith(".pdf") else "jpg"
                pages = extract_text_with_ocr(uploaded_file=uploaded_file, file_type=file_type)

                if pages:
                    text = "\n".join(pages)
                    summary = call_deepseek(
                        messages=[{"role": "user", "content": f"Summarize this document in clear, concise language:\n\n{text}"}],
                        temperature=0.3,
                        max_tokens=1024
                    )
                    st.session_state.go_buddy_summaries.append({
                        "filename": uploaded_file.name,
                        "summary": summary,
                        "text": text  # Store the full extracted text for chat
                    })
                    save_interaction(conn, "go_buddy", f"Document: {uploaded_file.name}", summary)

            progress_bar.progress((i + 1) / len(uploaded_files))

    if st.session_state.go_buddy_summaries:
        st.subheader("Document Summaries")
        for i, summary_data in enumerate(st.session_state.go_buddy_summaries):
            with st.expander(f"Summary for {summary_data['filename']}"):
                st.markdown(summary_data['summary'])

        if st.button("Download Merged Summaries"):
            doc = Document()
            doc.add_heading("Merged Document Summaries", level=1)
            doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            for summary_data in st.session_state.go_buddy_summaries:
                doc.add_heading(summary_data['filename'], level=2)
                doc.add_paragraph(summary_data['summary'])

            buf = io.BytesIO()
            doc.save(buf)
            st.download_button(
                "Download Merged Summaries",
                buf.getvalue(),
                file_name="merged_summaries.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # GO Buddy Chat functionality
        st.divider()
        st.subheader("Document Chat")

        # Display chat history
        for role, message in st.session_state.go_buddy_chat_memory:
            st.chat_message(role).write(message)

        # Chat input
        user_input = st.chat_input("Ask about these documents...")
        if user_input:
            # Add user message to chat
            st.session_state.go_buddy_chat_memory.append(("user", user_input))
            st.chat_message("user").write(user_input)

            # Build context from all documents
            context = "Documents:\n"
            for doc in st.session_state.go_buddy_summaries:
                context += f"\nDocument: {doc['filename']}\nSummary: {doc['summary']}\n"
            context += f"\nQuestion: {user_input}"

            # Call AI
            with st.spinner("Generating response..."):
                response = call_deepseek(
                    messages=[
                        {"role": "system", "content": "You are a document analysis expert. Answer questions based on the provided document summaries."},
                        {"role": "user", "content": context}
                    ],
                    temperature=0.3
                )

            # Add AI response to chat
            st.session_state.go_buddy_chat_memory.append(("assistant", response))
            st.chat_message("assistant").write(response)

            # Save interaction
            save_interaction(conn, "go_buddy_chat", user_input, response)

def admin_portal_ui(conn):
    """Admin portal with enhanced functionality"""
    st.header("🔒 Admin Portal")

    tab1, tab2, tab3, tab4, tab5 = st.tabs(["User Management", "Subscription Management", "Content Management", "Usage Analytics", "Prompt Management"])

    with tab1:
        st.subheader("User Accounts")
        cursor = conn.cursor()
        cursor.execute("""
            SELECT column_name
            FROM information_schema.columns
            WHERE table_name = 'users'
        """)
        columns = [column['column_name'] for column in cursor.fetchall()]

        select_columns = ["username", "role"]
        if "last_login" in columns:
            select_columns.append("last_login")
        if "location_id" in columns:
            select_columns.append("location_id")
        if "created_at" in columns:
            select_columns.append("created_at")

        query = f"SELECT {', '.join(select_columns)} FROM users"
        cursor.execute(query)
        users = cursor.fetchall()

        formatted_users = []
        for user in users:
            formatted_user = []
            for col in select_columns:
                value = user[col]
                if isinstance(value, str) and col in ['last_login', 'created_at']:
                    try:
                        value = datetime.strptime(value, "%Y-%m-%d %H:%M:%S").strftime("%Y-%m-%d %H:%M")
                    except:
                        pass
                formatted_user.append(value)
            formatted_users.append(formatted_user)

        user_df = pd.DataFrame(formatted_users, columns=select_columns)
        st.dataframe(user_df)

        with st.expander("Create New User"):
            with st.form("create_user_form"):
                new_user = st.text_input("Username")
                new_pass = st.text_input("Password", type="password")
                user_role = st.selectbox("Role", ["user", "admin"])
                location_id = st.text_input("Location ID")
                submitted = st.form_submit_button("Add User")

                if submitted:
                    if not new_user or not new_pass:
                        st.error("Username and password are required")
                    elif len(new_pass) < 8:
                        st.error("Password must be at least 8 characters")
                    else:
                        hashed = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt())
                        try:
                            cursor = conn.cursor()
                            if "location_id" in columns:
                                cursor.execute(
                                    "INSERT INTO users (username, password, role, location_id) VALUES (%s, %s, %s, %s)",
                                    (new_user, hashed, user_role, location_id)
                                )
                            else:
                                cursor.execute(
                                    "INSERT INTO users (username, password, role) VALUES (%s, %s, %s)",
                                    (new_user, hashed, user_role)
                                )
                            conn.commit()
                            st.success("User created successfully!")
                            time.sleep(1)
                            st.rerun()
                        except psycopg2.IntegrityError:
                            st.error("Username already exists")

    with tab2:
        st.subheader("Feature Access Control")
        cursor = conn.cursor()
        cursor.execute("SELECT username FROM users")
        users = cursor.fetchall()
        if not users:
            st.warning("No users found")
        else:
            selected_user = st.selectbox("Select User", [u['username'] for u in users])

            cursor.execute(
                "SELECT lease_analysis, deal_structuring, offer_generator FROM subscriptions WHERE username = %s",
                (selected_user,)
            )
            sub = cursor.fetchone()

            if not sub:
                cursor.execute(
                    "INSERT INTO subscriptions (username) VALUES (%s)",
                    (selected_user,)
                )
                conn.commit()
                sub = {'lease_analysis': 0, 'deal_structuring': 0, 'offer_generator': 0}

            with st.form("update_access_form"):
                col1, col2, col3 = st.columns(3)
                with col1:
                    lease_access = st.toggle("Lease Analysis", value=bool(sub['lease_analysis']))
                with col2:
                    deal_access = st.toggle("Deal Structuring", value=bool(sub['deal_structuring']))
                with col3:
                    offer_access = st.toggle("Offer Generator", value=bool(sub['offer_generator']))

                submitted = st.form_submit_button("Update Access")

                if submitted:
                    cursor.execute(
                        """UPDATE subscriptions
                        SET lease_analysis = %s, deal_structuring = %s, offer_generator = %s
                        WHERE username = %s""",
                        (int(lease_access), int(deal_access), int(offer_access), selected_user)
                    )
                    conn.commit()
                    st.success("Access updated successfully!")

    with tab3:
        st.subheader("Training Content")
        with st.expander("Upload Training Materials"):
            file_type = st.selectbox("Content Type", ["Document", "Video"])
            uploaded = st.file_uploader(
                f"Upload {file_type}",
                type=["pdf", "docx", "mp4"] if file_type == "Document" else ["mp4", "mov"]
            )
            description = st.text_area("Content Description")

            if uploaded and st.button("Upload"):
                save_dir = "training_content"
                os.makedirs(save_dir, exist_ok=True)
                file_path = os.path.join(save_dir, uploaded.name)

                with open(file_path, "wb") as f:
                    f.write(uploaded.getbuffer())

                meta_path = os.path.join(save_dir, f"{uploaded.name}.meta")
                with open(meta_path, "w") as f:
                    json.dump({
                        "uploaded_by": st.session_state.username,
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "description": description,
                        "type": file_type.lower()
                    }, f)

                st.success(f"{file_type} uploaded successfully!")

        st.subheader("Content Library")
        if os.path.exists("training_content"):
            files = os.listdir("training_content")
            content_files = [f for f in files if not f.endswith(".meta")]

            for file in content_files:
                meta_file = f"{file}.meta"
                if meta_file in files:
                    with open(os.path.join("training_content", meta_file)) as f:
                        meta = json.load(f)
                    st.markdown(f"**{file}**")
                    st.caption(f"Type: {meta['type']} | Uploaded by: {meta['uploaded_by']}")
                    st.caption(f"Description: {meta['description']}")
                    st.download_button(
                        f"Download {file}",
                        data=open(os.path.join("training_content", file), "rb").read(),
                        file_name=file
                    )
                    st.divider()

    with tab4:
        st.subheader("Usage Analytics")
        st.write("### Feature Usage")
        cursor = conn.cursor()
        cursor.execute(
            "SELECT feature, COUNT(*) as count FROM interactions GROUP BY feature"
        )
        usage = cursor.fetchall()
        if usage:
            fig = px.pie(
                names=[u['feature'] for u in usage],
                values=[u['count'] for u in usage],
                title="Feature Usage Distribution"
            )
            st.plotly_chart(fig)
        else:
            st.warning("No usage data available yet")

        st.write("### User Activity")
        cursor.execute(
            "SELECT username, COUNT(*) as interactions "
            "FROM interactions GROUP BY username ORDER BY interactions DESC LIMIT 10"
        )
        activity = cursor.fetchall()
        if activity:
            fig = px.bar(
                x=[a['username'] for a in activity],
                y=[a['interactions'] for a in activity],
                labels={"x": "User", "y": "Interactions"},
                title="Top Users by Activity"
            )
            st.plotly_chart(fig)
        else:
            st.warning("No user activity data available")

    with tab5:
        st.subheader("Prompt Management")

        # Select feature to edit
        feature = st.selectbox(
            "Select Feature",
            ["lease_analysis", "deal_strategy", "offer_generator"],
            format_func=lambda x: x.replace("_", " ").title()
        )

        # Get current prompts
        cursor = conn.cursor()
        cursor.execute(
            "SELECT system_prompt, user_prompt_template FROM prompts WHERE feature = %s",
            (feature,)
        )
        prompt = cursor.fetchone()

        if prompt:
            with st.form("prompt_update_form"):
                system_prompt = st.text_area(
                    "System Prompt",
                    value=prompt['system_prompt'],
                    height=150
                )
                user_prompt = st.text_area(
                    "User Prompt Template",
                    value=prompt['user_prompt_template'],
                    height=150
                )

                submitted = st.form_submit_button("Save Changes")

                if submitted:
                    cursor.execute(
                        "UPDATE prompts SET system_prompt = %s, user_prompt_template = %s, updated_at = CURRENT_TIMESTAMP WHERE feature = %s",
                        (system_prompt, user_prompt, feature)
                    )
                    conn.commit()
                    st.success("Prompts updated successfully!")
        else:
            st.warning("No prompts found for this feature")

def history_ui(conn):
    """Enhanced history view with better navigation"""
    st.header("🕒 Your History")

    if "username" not in st.session_state:
        st.warning("Please log in to view your history")
        return

    # If user has requested a full view, show it and bail out immediately
    if "current_interaction" in st.session_state:
        interaction = st.session_state.current_interaction
        st.subheader(f"Full Interaction – {interaction['timestamp']}")
        st.write(f"**Feature:** {interaction['feature']}")
        tabs = st.tabs(["Input", "Output"])
        with tabs[0]:
            st.text(interaction["input"])
        with tabs[1]:
            st.markdown(interaction["output"])

        if st.button("← Back to History"):
            del st.session_state.current_interaction
            st.rerun()
        return  # don't render the list below

    # Otherwise: render the list of past interactions
    cursor = conn.cursor()
    cursor.execute(
        "SELECT timestamp, feature, input_text, output_text "
        "FROM interactions WHERE username = %s ORDER BY timestamp DESC",
        (st.session_state.username,)
    )
    history = cursor.fetchall()

    if not history:
        st.info("No history found – your interactions will appear here")
        return

    for i, row in enumerate(history):
        ts, feature, inp, out = row['timestamp'], row['feature'], row['input_text'], row['output_text']
        with st.expander(f"{ts} • {feature}"):
            col1, col2 = st.columns(2)
            with col1:
                st.write("**Input**")
                st.text(inp[:500] + ("…" if len(inp) > 500 else ""))
            with col2:
                st.write("**Output**")
                st.text(out[:500] + ("…" if len(out) > 500 else ""))

            # single button per interaction
            if st.button(f"View Full Interaction #{i+1}", key=f"view_full_{i}"):
                st.session_state.current_interaction = {
                    "timestamp": ts,
                    "feature": feature,
                    "input": inp,
                    "output": out
                }
                st.rerun()

# def chatbot_ui(conn):
#     """Enhanced chatbot with better context handling"""
#     if not st.session_state.get("username"):
#         st.warning("Please log in to use the chatbot.")
#         return

#     # Initialize chat memory if not exists
#     if "chat_memory" not in st.session_state:
#         st.session_state["chat_memory"] = []

#     st.header("🤖 AI Chatbot")

#     # Clear chat button
#     if st.button("Clear Chat", key="clear_chat_button"):
#         st.session_state["chat_memory"] = []

#     st.markdown("Chat with the assistant based on your recent output.")

#     # Display past messages
#     for role, message in st.session_state["chat_memory"]:
#         st.chat_message(role).write(message)

#     # New user message
#     user_input = st.chat_input("Type your question...")
#     if user_input:
#         st.session_state["chat_memory"].append(("user", user_input))

#         # Build context from last 10 interactions
#         cursor = conn.cursor()
#         cursor.execute(
#             "SELECT feature, input_text, output_text FROM interactions WHERE username=%s ORDER BY timestamp DESC LIMIT 10",
#             (st.session_state.username,)
#         )
#         rows = cursor.fetchall()
#         context = "\n\n".join([f"Feature: {r['feature']}\nInput: {r['input_text']}\nOutput: {r['output_text']}" for r in rows])
#         prompt = f"Context:\n{context}\n\nQuestion:\n{user_input}"

#         # Call AI
#         if st.session_state.get("chat_model_choice", "Gemini") == "Gemini":
#             answer = call_gemini(conn, "chatbot", prompt)
#         elif st.session_state.get("chat_model_choice") == "Mistral":
#             messages = [
#                 {"role": "system", "content": "You are a helpful assistant using past interactions."},
#                 {"role": "user", "content": prompt}
#             ]
#             answer = call_deepseek(messages)
#         else:  # DeepSeek
#             messages = [
#                 {"role": "system", "content": "You are an AI assistant answering questions based on the user's context."},
#                 {"role": "user", "content": prompt}
#             ]
#             answer = call_deepseek(messages)

#         # Append and display bot response
#         st.session_state["chat_memory"].append(("assistant", answer))
#         st.chat_message("assistant").write(answer)

#         # Save interaction
#         save_interaction(conn, "chatbot", user_input, answer)

def register_ui(conn):
    """Enhanced registration with better validation"""
    st.title("📝 Register")
    st.markdown("Create a new account.")

    with st.form("register_form"):
        username = st.text_input("Username", placeholder="Enter your username")
        password = st.text_input("Password", type="password", placeholder="Enter your password")
        confirm_password = st.text_input("Confirm Password", type="password", placeholder="Confirm your password")
        location_id = st.text_input("Location ID (optional)", placeholder="Enter location ID if applicable")
        submitted = st.form_submit_button("Register")

        if submitted:
            # Input validation
            if not username or not password or not confirm_password:
                st.error("All required fields must be filled.")
                return
            if len(password) < 8:
                st.error("Password must be at least 8 characters long.")
                return
            if password != confirm_password:
                st.error("Passwords do not match.")
                return
            if not re.match(r"^[a-zA-Z0-9_]+$", username):
                st.error("Username can only contain letters, numbers, and underscores.")
                return

            # Check for existing username
            cursor = conn.cursor()
            cursor.execute("SELECT username FROM users WHERE username = %s", (username,))
            if cursor.fetchone():
                st.error("Username already exists. Please choose a different username.")
                return

            # Register user
            try:
                hashed = bcrypt.hashpw(password.encode(), bcrypt.gensalt())
                cursor.execute(
                    "INSERT INTO users (username, password, role, location_id, created_at) VALUES (%s, %s, %s, %s, %s)",
                    (username, hashed, "user", location_id or None, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                )
                cursor.execute(
                    "INSERT INTO subscriptions (username, lease_analysis, deal_structuring, offer_generator) VALUES (%s, %s, %s, %s)",
                    (username, 0, 0, 0)
                )
                conn.commit()
                st.success("Registration successful! You can now log in.")
                logging.info(f"User {username} registered successfully")
            except psycopg2.Error as e:
                st.error(f"Registration failed: {e}")
                logging.error(f"Registration failed for {username}: {e}")

def login_ui(conn):
    """Enhanced login with better error handling"""
    st.title("🔐 Property Deals AI")
    st.markdown("Access or create your account to use the platform.")

    # Create tabs for Login and Register
    login_tab, register_tab = st.tabs(["Login", "Register"])

    with login_tab:
        st.markdown("### Log In")
        with st.form("login_form"):
            username = st.text_input("Username", key="login_username")
            password = st.text_input("Password", type="password", key="login_password")
            submitted = st.form_submit_button("Login")

            if submitted:
                if not username or not password:
                    st.error("Please provide both username and password.")
                    return

                cursor = conn.cursor()
                cursor.execute(
                    "SELECT password, role, location_id FROM users WHERE username = %s",
                    (username,)
                )
                user = cursor.fetchone()

                if user and verify_password(user['password'], password):
                    st.session_state.logged_in = True
                    st.session_state.username = username
                    st.session_state.role = user['role']
                    st.session_state.location_id = user['location_id'] if user['location_id'] else None

                    # Update last login timestamp
                    cursor.execute(
                        "UPDATE users SET last_login = %s WHERE username = %s",
                        (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), username)
                    )
                    conn.commit()

                    st.success(f"Welcome, {username}!")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("Invalid username or password.")

    with register_tab:
        register_ui(conn)

def main():
    """Main application function with improved initialization"""
    # Configure page
    st.set_page_config(
        page_title="Property Deals AI",
        page_icon="🏠",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Apply brand styling
    st.markdown(
        """
        <style>
        /* General sidebar styling */
        .sidebar .sidebar-content {
            padding: 20px;
            background-color: #1a1a2e;
            color: #ffffff;
        }

        /* Welcome message styling */
        .sidebar .welcome {
            font-size: 1.5em;
            font-weight: 600;
            margin-bottom: 10px;
            color: #ffffff;
        }

        /* Location ID styling */
        .sidebar .location {
            font-size: 0.9em;
            color: #b0b0b0;
            margin-bottom: 20px;
        }

        /* Navigation header styling */
        .sidebar .navigation-header {
            font-size: 1.2em;
            font-weight: 500;
            margin-bottom: 15px;
            color: #2E86AB;
            border-bottom: 1px solid #2E86AB;
            padding-bottom: 5px;
        }

        /* Button styling */
        .sidebar .stButton>button {
            display: block;
            width: 100%;
            height: 50px;
            background-color: #F18F01;
            color: white;
            border: none;
            border-radius: 8px;
            padding: 0 10px;
            margin-bottom: 10px;
            font-size: 1em;
            font-weight: 500;
            text-align: center;
            line-height: 50px;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            transition: all 0.3s ease;
        }

        .sidebar .stButton>button:hover {
            background-color: #d67a00;
            transform: translateY(-2px);
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
        }

        .sidebar .stButton>button:last-child {
            margin-bottom: 0;
        }

        /* Main content styling */
        .main .block-container {
            padding: 20px;
        }

        /* Chat history boxes */
        .st-chat-message {
            border-radius: 8px !important;
            padding: 12px !important;
            margin-bottom: 8px !important;
        }
        .st-chat-message.user {
            background-color: #e8f1fb !important;
        }
        .st-chat-message.assistant {
            background-color: #f7f7f7 !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Initialize database with loading indicator
    if "_initialized" not in st.session_state:
        with st.spinner("Initializing application..."):
            try:
                init_db()
                st.session_state._initialized = True
            except Exception as e:
                st.error(f"Failed to initialize application: {e}")
                return

    # Ensure login state
    if "logged_in" not in st.session_state:
        st.session_state.update({
            "logged_in": False,
            "username": None,
            "role": None,
            "subscription": {
                "lease_analysis": False,
                "deal_structuring": False,
                "offer_generator": False
            }
        })

    # Authentication flow
    if not st.session_state.logged_in:
        with get_db_connection() as conn:
            login_ui(conn)
        return

    # Get user's subscription status
    if st.session_state.logged_in:
        try:
            with get_db_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT lease_analysis, deal_structuring, offer_generator FROM subscriptions WHERE username = %s",
                    (st.session_state.username,)
                )
                sub = cursor.fetchone()

                if not sub and st.session_state.role != "admin":
                    cursor.execute(
                        "INSERT INTO subscriptions (username) VALUES (%s)",
                        (st.session_state.username,)
                    )
                    conn.commit()
                    sub = {'lease_analysis': 0, 'deal_structuring': 0, 'offer_generator': 0}
                elif st.session_state.role == "admin":
                    sub = {'lease_analysis': 1, 'deal_structuring': 1, 'offer_generator': 1}

                st.session_state.subscription = {
                    "lease_analysis": bool(sub['lease_analysis']),
                    "deal_structuring": bool(sub['deal_structuring']),
                    "offer_generator": bool(sub['offer_generator'])
                }
        except psycopg2.Error as e:
            st.error(f"Error fetching subscription status: {e}")
            return

    # Sidebar navigation
    st.sidebar.markdown('<div class="sidebar-content">', unsafe_allow_html=True)
    st.sidebar.markdown(f'<div class="welcome">Welcome, {st.session_state.username}</div>', unsafe_allow_html=True)
    st.sidebar.markdown(f'<div class="location">Location ID: {st.session_state.get("location_id", "None")}</div>', unsafe_allow_html=True)

    # Available features based on subscription
    features = []
    if st.session_state.subscription.get("lease_analysis") or st.session_state.role == "admin":
        features.append("LeaseBrief Buddy")
    if st.session_state.subscription.get("deal_structuring") or st.session_state.role == "admin":
        features.append("Property Guru")
    if st.session_state.subscription.get("offer_generator") or st.session_state.role == "admin":
        features.append("Offer Buddy")
    features.append("Auction Buddy")
    features.append("History")
    if st.session_state.role == "admin":
        features.insert(-1, "Admin Portal")

    # Initialize selected feature
    if "selected_feature" not in st.session_state:
        st.session_state.selected_feature = features[0] if features else None

    # Navigation buttons
    st.sidebar.markdown('<div class="navigation-header">Navigation</div>', unsafe_allow_html=True)
    for feature in features:
        if st.sidebar.button(feature, key=f"nav_{feature.replace(' ', '_')}"):
            st.session_state.selected_feature = feature
            st.rerun()
    st.sidebar.markdown('</div>', unsafe_allow_html=True)

    # Main content area
    try:
        with get_db_connection() as conn:
            selected = st.session_state.selected_feature
            if selected == "LeaseBrief Buddy":
                lease_summarization_ui(conn)
            elif selected == "Property Guru":
                deal_structuring_ui(conn)
            elif selected == "Offer Buddy":
                offer_generator_ui(conn)
            elif selected == "Auction Buddy":
                go_buddy_ui(conn)
            elif selected == "History":
                history_ui(conn)
            elif selected == "Admin Portal" and st.session_state.role == "admin":
                admin_portal_ui(conn)
            else:
                st.error("Access Denied")
    except Exception as e:
        st.error(f"Application error: {e}")

    # Chatbot section
    # st.divider()
    # try:
    #     with get_db_connection() as conn:
    #         chatbot_ui(conn)
    # except Exception as e:
    #     st.error(f"Chatbot error: {e}")

    # Logout button
    st.sidebar.markdown('<div class="sidebar-content">', unsafe_allow_html=True)
    st.sidebar.divider()
    if st.sidebar.button("Logout"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()
    st.sidebar.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()